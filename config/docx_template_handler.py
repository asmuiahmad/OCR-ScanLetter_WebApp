"""
DOCX Template Handler
Handles Word document template filling and PDF conversion
"""

import os
import tempfile
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Inches
import qrcode
from io import BytesIO
import base64

class DocxTemplateHandler:
    def __init__(self):
        self.template_path = 'static/assets/templates/form_permintaan_cuti.docx'
        self.output_folder = 'static/pdf_cuti'
        self.signatures_folder = 'static/signatures'
        
        # Create output directories
        os.makedirs(self.output_folder, exist_ok=True)
        os.makedirs(self.signatures_folder, exist_ok=True)
    
    def create_qr_code(self, cuti_data, signature_hash):
        """Create QR code for digital signature"""
        try:
            # Data untuk QR code
            qr_data = f"PERSETUJUAN CUTI\n"
            qr_data += f"ID: {cuti_data.id_cuti}\n"
            qr_data += f"Nama: {cuti_data.nama}\n"
            qr_data += f"NIP: {cuti_data.nip}\n"
            qr_data += f"Jenis: {cuti_data.jenis_cuti}\n"
            qr_data += f"Periode: {cuti_data.tanggal_cuti.strftime('%Y-%m-%d')} s/d {cuti_data.sampai_cuti.strftime('%Y-%m-%d')}\n"
            qr_data += f"Tanggal: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            qr_data += f"Hash: {signature_hash}\n"
            qr_data += f"Verifikasi: https://your-domain.com/verify/{signature_hash}"
            
            # Generate QR code
            qr = qrcode.QRCode(
                version=1,
                error_correction=qrcode.constants.ERROR_CORRECT_L,
                box_size=10,
                border=4,
            )
            qr.add_data(qr_data)
            qr.make(fit=True)
            
            # Create QR code image
            qr_img = qr.make_image(fill_color="black", back_color="white")
            
            # Save QR code file
            qr_filename = f"qr_cuti_{cuti_data.id_cuti}_{signature_hash}.png"
            qr_path = os.path.join(self.signatures_folder, qr_filename)
            qr_img.save(qr_path)
            
            return qr_path
            
        except Exception as e:
            print(f"Error creating QR code: {str(e)}")
            return None
    
    def generate_signature_hash(self, cuti_data):
        """Generate unique hash for digital signature"""
        import hashlib
        signature_string = f"{cuti_data.id_cuti}_{cuti_data.nama}_{cuti_data.nip}_{datetime.now().isoformat()}"
        return hashlib.sha256(signature_string.encode()).hexdigest()[:16]
    
    def format_jenis_cuti(self, jenis_cuti):
        """Convert jenis_cuti code to readable format"""
        jenis_map = {
            'c_tahun': 'Cuti Tahunan',
            'c_besar': 'Cuti Besar',
            'c_sakit': 'Cuti Sakit',
            'c_lahir': 'Cuti Melahirkan',
            'c_penting': 'Cuti Alasan Penting',
            'c_luarnegara': 'Cuti Luar Negara'
        }
        return jenis_map.get(jenis_cuti, jenis_cuti)
    
    def format_date_indonesian(self, date_obj):
        """Format date to Indonesian format"""
        if not date_obj:
            return ""
        
        months = [
            'Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni',
            'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember'
        ]
        
        return f"{date_obj.day} {months[date_obj.month - 1]} {date_obj.year}"
    
    def replace_placeholders_in_docx(self, doc, cuti_data, signature_hash):
        """Replace placeholders in Word document with actual data"""
        try:
            # Create replacement dictionary with the format used in template
            replacements = {
                '«nama»': cuti_data.nama,
                '«nip»': cuti_data.nip,
                '«jabatan»': cuti_data.jabatan,
                '«gol_ruang»': cuti_data.gol_ruang,
                '«unit_kerja»': cuti_data.unit_kerja,
                '«masa_kerja»': cuti_data.masa_kerja,
                '«alamat»': cuti_data.alamat,
                '«telp»': cuti_data.telp,
                '«no_suratmasuk»': cuti_data.no_suratmasuk,
                '«tgl_lengkap_ajuan_cuti»': self.format_date_indonesian(cuti_data.tgl_ajuan_cuti),
                '«bulan_ajuan_cuti»': str(cuti_data.tgl_ajuan_cuti.month).zfill(2),
                '«tahun_ajuan_cuti»': str(cuti_data.tgl_ajuan_cuti.year),
                '«alasan_cuti»': cuti_data.alasan_cuti,
                '«lama_cuti»': cuti_data.lama_cuti.replace(' hari', ''),  # Remove 'hari' as template has it
                '«tanggal_cuti»': self.format_date_indonesian(cuti_data.tanggal_cuti),
                '«sampai_cuti»': self.format_date_indonesian(cuti_data.sampai_cuti),
                # Jenis cuti checkboxes - mark the selected one
                '«c_tahun»': '✓' if cuti_data.jenis_cuti == 'c_tahun' else '',
                '«c_besar»': '✓' if cuti_data.jenis_cuti == 'c_besar' else '',
                '«c_sakit»': '✓' if cuti_data.jenis_cuti == 'c_sakit' else '',
                '«c_lahir»': '✓' if cuti_data.jenis_cuti == 'c_lahir' else '',
                '«c_penting»': '✓' if cuti_data.jenis_cuti == 'c_penting' else '',
                '«c_luarnegara»': '✓' if cuti_data.jenis_cuti == 'c_luarnegara' else '',
            }
            
            # Replace in paragraphs
            for paragraph in doc.paragraphs:
                for placeholder, value in replacements.items():
                    if placeholder in paragraph.text:
                        # Handle runs to preserve formatting
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, str(value))
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, value in replacements.items():
                                if placeholder in paragraph.text:
                                    # Handle runs to preserve formatting
                                    for run in paragraph.runs:
                                        if placeholder in run.text:
                                            run.text = run.text.replace(placeholder, str(value))
            
            return True
            
        except Exception as e:
            print(f"Error replacing placeholders: {str(e)}")
            return False
    
    def add_qr_code_to_docx(self, doc, qr_path):
        """Add QR code image to Word document"""
        try:
            if os.path.exists(qr_path):
                # Find paragraph with QR placeholder or add to end
                qr_paragraph = None
                for paragraph in doc.paragraphs:
                    if '{{QR_CODE}}' in paragraph.text:
                        qr_paragraph = paragraph
                        break
                
                if qr_paragraph:
                    # Clear the placeholder text
                    qr_paragraph.clear()
                    # Add QR code image
                    run = qr_paragraph.runs[0] if qr_paragraph.runs else qr_paragraph.add_run()
                    run.add_picture(qr_path, width=Inches(1.5))
                else:
                    # Add QR code at the end
                    paragraph = doc.add_paragraph()
                    paragraph.add_run().add_picture(qr_path, width=Inches(1.5))
                
                return True
        except Exception as e:
            print(f"Error adding QR code: {str(e)}")
            return False
    
    def docx_to_pdf_reportlab(self, docx_path, pdf_path):
        """Convert DOCX to PDF using reportlab (manual conversion)"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER, TA_LEFT
            
            # Read the docx file
            doc = Document(docx_path)
            
            # Create PDF
            pdf_doc = SimpleDocTemplate(pdf_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Custom styles
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=6,
                fontName='Helvetica'
            )
            
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=14,
                spaceAfter=12,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            
            # Convert paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    if any(keyword in paragraph.text.upper() for keyword in ['FORMULIR', 'PERMINTAAN', 'CUTI']):
                        story.append(Paragraph(paragraph.text, title_style))
                    else:
                        story.append(Paragraph(paragraph.text, normal_style))
                    story.append(Spacer(1, 6))
            
            # Convert tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = ""
                        for paragraph in cell.paragraphs:
                            cell_text += paragraph.text + " "
                        row_data.append(cell_text.strip())
                    table_data.append(row_data)
                
                if table_data:
                    pdf_table = Table(table_data)
                    pdf_table.setStyle(TableStyle([
                        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 0), (-1, -1), 9),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('LEFTPADDING', (0, 0), (-1, -1), 6),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                        ('TOPPADDING', (0, 0), (-1, -1), 6),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ]))
                    story.append(pdf_table)
                    story.append(Spacer(1, 12))
            
            # Build PDF
            pdf_doc.build(story)
            return True
            
        except Exception as e:
            print(f"Error converting DOCX to PDF: {str(e)}")
            return False
    
    def fill_template_and_generate_pdf(self, cuti_data):
        """Main function to fill template and generate PDF"""
        try:
            # Check if template exists
            if not os.path.exists(self.template_path):
                raise FileNotFoundError(f"Template not found: {self.template_path}")
            
            # Generate signature hash
            signature_hash = self.generate_signature_hash(cuti_data)
            
            # Create QR code
            qr_path = self.create_qr_code(cuti_data, signature_hash)
            
            # Load template
            doc = Document(self.template_path)
            
            # Replace placeholders
            if not self.replace_placeholders_in_docx(doc, cuti_data, signature_hash):
                raise Exception("Failed to replace placeholders")
            
            # Add QR code if created
            if qr_path:
                self.add_qr_code_to_docx(doc, qr_path)
            
            # Save filled docx temporarily
            temp_docx = os.path.join(self.output_folder, f"temp_cuti_{cuti_data.id_cuti}_{signature_hash}.docx")
            doc.save(temp_docx)
            
            # Generate PDF filename
            pdf_filename = f"surat_cuti_{cuti_data.id_cuti}_{signature_hash}.pdf"
            pdf_path = os.path.join(self.output_folder, pdf_filename)
            
            # Convert to PDF
            if self.docx_to_pdf_reportlab(temp_docx, pdf_path):
                # Clean up temporary docx
                if os.path.exists(temp_docx):
                    os.remove(temp_docx)
                
                return {
                    'success': True,
                    'pdf_path': pdf_path,
                    'qr_path': qr_path,
                    'signature_hash': signature_hash
                }
            else:
                raise Exception("Failed to convert DOCX to PDF")
                
        except Exception as e:
            print(f"Error in fill_template_and_generate_pdf: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }